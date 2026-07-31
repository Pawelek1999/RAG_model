"""Document loading layer converting supported files into LangChain documents."""

import logging
import os
import re
from pathlib import Path
from typing import Callable

import pandas as pd
from docx import Document as DocxDocument
from langchain_core.documents import Document
from pypdf import PdfReader

from backend.tools.Excel_tests import ExcelJsonFormatter, ExcelWorkbook, JsonFormattingOptions


logger = logging.getLogger(__name__)


class DocumentLoader:
    """Loads supported file types and normalizes them into document objects."""

    _TEST_SEQUENCE_PATTERN = re.compile(
        r"^(?P<prefix>[A-Za-z0-9]+_)?(?P<test_number>\d{5})\.(?P<step_number>\d{3})$"
    )
    _BUG_NUMBER_PATTERN = re.compile(r"\bBUG(?:\s*NB)?\s*[:#-]?\s*(?P<bug_number>\d+)\b", re.IGNORECASE)

    def __init__(self, xlsx_mode: str | None = None) -> None:
        """Initializes loader mappings and spreadsheet loading strategy.

        Args:
            xlsx_mode: Spreadsheet mode: auto, standard, or test-oriented.
        """
        self._xlsx_mode = self._normalize_xlsx_mode(
            xlsx_mode or os.getenv("XLSX_LOADER_MODE", "auto")
        )
        self._loaders: dict[str, Callable[[Path], list[Document]]] = {
            ".docx": self._load_docx,
            ".pdf": self._load_pdf,
            ".txt": self._load_text,
            ".md": self._load_text,
            ".xlsx": self._load_xlsx,
        }

    def load(self, file_path: str | Path) -> list[Document]:
        """Loads one file and returns normalized LangChain documents.

        Args:
            file_path: Path to a supported document file.

        Returns:
            Documents extracted from the file.

        Raises:
            FileNotFoundError: When the path does not exist.
            ValueError: When the path is invalid or extension unsupported.
        """
        path = Path(file_path)
        self._validate_file(path)

        extension = path.suffix.lower()
        logger.info("loader-load-start path=%s extension=%s", path, extension)
        loader = self._loaders.get(extension)

        if loader is None:
            supported = ", ".join(sorted(self._loaders))
            raise ValueError(
                f"Nieobslugiwany format pliku: {extension}. "
                f"Obslugiwane formaty: {supported}"
            )

        try:
            documents = loader(path)
            logger.info(
                "loader-load-end path=%s extension=%s documents_count=%s",
                path,
                extension,
                len(documents),
            )
            return documents
        except Exception:
            logger.exception("loader-load-error path=%s extension=%s", path, extension)
            raise

    def supported_extensions(self) -> list[str]:
        """Returns supported file extensions.

        Returns:
            Sorted list of recognized extensions.
        """
        return sorted(self._loaders)

    def _validate_file(self, path: Path) -> None:
        # Sprawdza, czy podana sciezka istnieje i wskazuje na zwykly plik.
        if not path.exists():
            raise FileNotFoundError(f"Plik nie istnieje: {path}")

        if not path.is_file():
            raise ValueError(f"Podana sciezka nie jest plikiem: {path}")

    def _load_docx(self, path: Path) -> list[Document]:
        # Wczytuje tekst z pliku DOCX, laczac niepuste akapity w jeden dokument.
        docx = DocxDocument(path)
        paragraphs = [
            paragraph.text.strip()
            for paragraph in docx.paragraphs
            if paragraph.text.strip()
        ]
        text = "\n".join(paragraphs)

        return [
            Document(
                page_content=text,
                metadata=self._base_metadata(path, file_type="docx"),
            )
        ]

    def _load_pdf(self, path: Path) -> list[Document]:
        # Wczytuje PDF i tworzy osobny Document dla kazdej strony z tekstem.
        reader = PdfReader(path)
        documents: list[Document] = []

        for page_number, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            text = text.strip()

            if not text:
                continue

            metadata = self._base_metadata(path, file_type="pdf")
            metadata["page"] = page_number
            metadata["total_pages"] = len(reader.pages)

            documents.append(Document(page_content=text, metadata=metadata))

        return documents

    def _load_text(self, path: Path) -> list[Document]:
        # Wczytuje zwykly plik tekstowy, na przyklad TXT albo Markdown.
        text = path.read_text(encoding="utf-8").strip()

        return [
            Document(
                page_content=text,
                metadata=self._base_metadata(path, file_type=path.suffix.lower().lstrip(".")),
            )
        ]

    def _load_xlsx(self, path: Path) -> list[Document]:
        # Wczytuje plik XLSX i wybiera odpowiedni tryb dla workbooka.
        mode = self._resolve_xlsx_mode(path)
        logger.info("loader-xlsx-read-start path=%s mode=%s", path, mode)

        if mode == "test-oriented":
            return self._load_xlsx_test_oriented(path)

        return self._load_xlsx_standard(path)

    def _load_xlsx_standard(self, path: Path) -> list[Document]:
        # Zachowuje dotychczasowy przeplyw dla zwyklych arkuszy XLSX.
        logger.info("loader-xlsx-standard-read path=%s engine=openpyxl", path)
        sheets = pd.read_excel(path, sheet_name=None, engine="openpyxl")
        documents: list[Document] = []

        for sheet_name, dataframe in sheets.items():
            dataframe = dataframe.dropna(how="all").dropna(axis=1, how="all")

            if dataframe.empty:
                continue

            text = dataframe.to_csv(index=False).strip()
            metadata = self._base_metadata(path, file_type="xlsx")
            metadata["sheet_name"] = str(sheet_name)

            documents.append(Document(page_content=text, metadata=metadata))

        logger.info(
            "loader-xlsx-standard-end path=%s sheets_count=%s documents_count=%s",
            path,
            len(sheets),
            len(documents),
        )
        return documents

    def _load_xlsx_test_oriented(self, path: Path) -> list[Document]:
        # Zamienia wiersze testowe na semantyczne Document przed chunkowaniem.
        formatter = ExcelJsonFormatter(options=JsonFormattingOptions())
        payload = formatter.format_workbook(workbook_path=path)
        documents: list[Document] = []
        rows_total = 0
        rows_rejected = 0

        for sheet_payload in payload.get("sheets", []):
            sheet_name = str(sheet_payload.get("sheet_name") or "")
            for row in sheet_payload.get("rows", []):
                rows_total += 1

                if self._is_test_header_row(row):
                    rows_rejected += 1
                    continue

                if row.get("skip_from_business_flow"):
                    rows_rejected += 1
                    continue

                if row.get("row_type") not in {"test_step", "test_header"}:
                    rows_rejected += 1
                    continue

                page_content = self._build_test_row_content(row)
                if not page_content.strip():
                    rows_rejected += 1
                    continue

                metadata = self._base_metadata(path, file_type="xlsx")
                metadata.update(self._build_test_row_metadata(sheet_name=sheet_name, row=row))

                documents.append(Document(page_content=page_content, metadata=metadata))

        logger.info(
            "loader-xlsx-test-oriented-end path=%s rows_total=%s rows_rejected=%s documents_count=%s",
            path,
            rows_total,
            rows_rejected,
            len(documents),
        )
        return documents

    def _is_test_header_row(self, row: dict[str, object]) -> bool:
        # Chroni indeks przed wczytaniem wiersza naglowkowego jako testu.
        row_index = row.get("row_index")
        if row_index != 1:
            return False

        importance = str(row.get("importance") or "").strip().casefold()
        procedure = str(row.get("procedure") or "").strip().casefold()
        expected_result = str(row.get("expected_result") or "").strip().casefold()

        return importance in {"importance", "nb d'importance", "nb d’importance"} or (
            procedure == "procedure" and expected_result in {"expected result", "expected_result"}
        )

    def _build_test_row_content(self, row: dict[str, object]) -> str:
        # Buduje semantyczny opis jednego wiersza testowego.
        display_fields = row.get("display_fields")
        if isinstance(display_fields, list) and display_fields:
            field_labels = [
                (
                    str(field.get("label") or "").strip(),
                    field.get("value"),
                )
                for field in display_fields
                if isinstance(field, dict)
            ]
        else:
            field_labels = [
                ("Test sequence number", row.get("test_sequence_number")),
                ("Revision", row.get("revision")),
                ("Procedure", row.get("procedure")),
                ("Expected result", row.get("expected_result")),
                ("Observed result", row.get("observed_result")),
                ("Conclusion", row.get("conclusion")),
                ("Comment", row.get("comment")),
                ("Status", row.get("status")),
                ("Anomaly", row.get("anomaly")),
            ]

        lines: list[str] = []
        for label, value in field_labels:
            if not label:
                continue
            text = str(value or "").strip()
            if not text:
                continue
            lines.append(f"{label}: {text}")

        return "\n".join(lines)

    def _build_test_row_metadata(self, sheet_name: str, row: dict[str, object]) -> dict[str, object]:
        # Przenosi metadane wiersza testowego do warstwy Document.
        test_sequence_number = row.get("test_sequence_number")
        sheet_id, test_number, step_number = self._extract_test_sequence_parts(test_sequence_number)
        observed_result = row.get("observed_result")
        conclusion = row.get("conclusion")
        bug_number = self._extract_bug_number(observed_result, conclusion)

        return {
            "sheet_name": sheet_name,
            "row_index": row.get("row_index"),
            "row_type": row.get("row_type"),
            "status": row.get("status"),
            "anomaly": row.get("anomaly"),
            "skip_from_business_flow": row.get("skip_from_business_flow"),
            "test_sequence_number": test_sequence_number,
            "sheet_id": sheet_id,
            "test_number": test_number,
            "step_number": step_number,
            "procedure": row.get("procedure"),
            "expected_result": row.get("expected_result"),
            "observed_result": observed_result,
            "conclusion": conclusion,
            "bug_number": bug_number,
            "revision": row.get("revision"),
        }

    def _extract_test_sequence_parts(
        self,
        raw_sequence: object,
    ) -> tuple[str | None, str | None, str | None]:
        # Normalizuje sheet_id, numer testu i numer kroku z pola test_sequence_number.
        sequence = str(raw_sequence or "").strip()
        if not sequence:
            return None, None, None

        match = self._TEST_SEQUENCE_PATTERN.match(sequence)
        if not match:
            return None, None, None

        prefix = str(match.group("prefix") or "").strip()
        sheet_id = prefix[:-1] if prefix.endswith("_") else prefix
        sheet_id = sheet_id or None

        return sheet_id, match.group("test_number"), match.group("step_number")

    def _extract_bug_number(self, *values: object) -> str | None:
        for value in values:
            text = str(value or "").strip()
            if not text:
                continue

            match = self._BUG_NUMBER_PATTERN.search(text)
            if match:
                return str(match.group("bug_number") or "").strip() or None

        return None

    def _resolve_xlsx_mode(self, path: Path) -> str:
        # Umozliwia jawny wybor trybu oraz automatyczne wykrycie workbooka testowego.
        if self._xlsx_mode != "auto":
            return self._xlsx_mode

        if self._looks_like_test_workbook(path):
            return "test-oriented"

        return "standard"

    def _looks_like_test_workbook(self, path: Path) -> bool:
        # Wykrywa workbook testowy po naglowku arkusza.
        workbook = ExcelWorkbook.load(path)

        for sheet in workbook.get_all_sheets():
            first_cell = sheet.get_cell(row=1, column=1).as_string().strip().casefold()
            if first_cell in {"importance", "nb d'importance", "nb d’importance"}:
                logger.debug("loader-xlsx-test-detected path=%s sheet=%s", path, sheet.name)
                return True

        return False

    def _normalize_xlsx_mode(self, mode: str) -> str:
        normalized = mode.strip().casefold().replace("_", "-")

        if normalized in {"auto", "standard", "test-oriented", "test"}:
            return "test-oriented" if normalized == "test" else normalized

        logger.warning("loader-xlsx-mode-unknown mode=%s fallback=auto", mode)
        return "auto"

    def _base_metadata(self, path: Path, file_type: str) -> dict[str, str | int | bool | None]:
        # Buduje wspolne metadane dodawane do kazdego wczytanego dokumentu.
        return {
            "source": str(path),
            "file_name": path.name,
            "file_type": file_type,
        }
